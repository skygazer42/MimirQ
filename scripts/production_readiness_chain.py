#!/usr/bin/env python3
"""Run a real production-readiness RAG/KG ingestion chain.

This is intentionally an integration runner, not a mocked unit test. It creates
an isolated dataset, uploads a multi-format corpus, waits for backend processing,
checks chunking/KG/RAG/chat behavior, and writes machine-readable evidence.
"""

import argparse
import csv
import json
import os
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from urllib.parse import urlencode

import requests
from docx import Document
from openpyxl import Workbook

TENANT_ID = "00000000-0000-0000-0000-000000000000"
USER_ID = "production-readiness"
UTC = timezone.utc


@dataclass(frozen=True)
class Source:
    filename: str
    url: str
    kind: str


@dataclass
class Evidence:
    started_at: str
    base_url: str
    tenant_id: str
    user_id: str
    corpus_dir: str
    output_dir: str
    settings: dict[str, Any] = field(default_factory=dict)
    dataset: dict[str, Any] = field(default_factory=dict)
    downloads: list[dict[str, Any]] = field(default_factory=list)
    generated_files: list[dict[str, Any]] = field(default_factory=list)
    uploads: list[dict[str, Any]] = field(default_factory=list)
    documents: list[dict[str, Any]] = field(default_factory=list)
    chunk_previews: list[dict[str, Any]] = field(default_factory=list)
    kg: dict[str, Any] = field(default_factory=dict)
    retrieval_warmups: list[dict[str, Any]] = field(default_factory=list)
    retrieval: list[dict[str, Any]] = field(default_factory=list)
    chat: list[dict[str, Any]] = field(default_factory=list)
    default_chat: list[dict[str, Any]] = field(default_factory=list)
    provider_health: dict[str, Any] = field(default_factory=dict)
    checks: list[dict[str, Any]] = field(default_factory=list)
    failures: list[str] = field(default_factory=list)

    def check(self, name: str, ok: bool, **fields: Any) -> None:
        row = {"name": name, "ok": bool(ok), **fields}
        self.checks.append(row)
        if not ok:
            self.failures.append(name + (f": {fields.get('reason')}" if fields.get("reason") else ""))


SOURCES: list[Source] = [
    Source("rfc9000-quic.txt", "https://www.rfc-editor.org/rfc/rfc9000.txt", "txt"),
    Source("rfc9110-http-semantics.txt", "https://www.rfc-editor.org/rfc/rfc9110.txt", "txt"),
    Source("rfc9000-quic.pdf", "https://www.rfc-editor.org/rfc/rfc9000.pdf", "pdf"),
    Source("rfc9110-http-semantics.pdf", "https://www.rfc-editor.org/rfc/rfc9110.pdf", "pdf"),
    Source("fastapi-readme.md", "https://raw.githubusercontent.com/fastapi/fastapi/master/README.md", "md"),
    Source("httpx-readme.md", "https://raw.githubusercontent.com/encode/httpx/master/README.md", "md"),
    Source("flask-readme.md", "https://raw.githubusercontent.com/pallets/flask/main/README.md", "md"),
    Source("python-asyncio.html", "https://docs.python.org/3/library/asyncio.html", "html"),
    Source("wcag21.html", "https://www.w3.org/TR/WCAG21/", "html"),
    Source("iris.csv", "https://raw.githubusercontent.com/mwaskom/seaborn-data/master/iris.csv", "csv"),
]


def now_id() -> str:
    return datetime.now(UTC).strftime("%Y%m%d-%H%M%S")


def json_dumps(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False, indent=2, default=str)


def fail(message: str) -> None:
    raise RuntimeError(message)


def load_local_env_value(key: str, *, env_path: Path = Path(".env")) -> str:
    """Read one local .env value without adding a dotenv dependency."""
    value = os.environ.get(key)
    if value:
        return value
    if not env_path.exists():
        return ""
    prefix = f"{key}="
    for raw_line in env_path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or not line.startswith(prefix):
            continue
        out = line[len(prefix) :].strip()
        if (out.startswith('"') and out.endswith('"')) or (out.startswith("'") and out.endswith("'")):
            out = out[1:-1]
        return out
    return ""


def load_llm_probe_api_key(api_base: str) -> tuple[str, str]:
    base = str(api_base or "").casefold()
    if "siliconflow" in base:
        return load_local_env_value("SILICONFLOW_API_KEY"), "SILICONFLOW_API_KEY"
    if "api.openai.com" in base:
        return load_local_env_value("OPENAI_API_KEY"), "OPENAI_API_KEY"
    if "dashscope" in base or "aliyuncs" in base:
        return load_local_env_value("LLM_API_KEY"), "LLM_API_KEY"
    return load_local_env_value("LLM_API_KEY"), "LLM_API_KEY"


class Api:
    def __init__(self, base_url: str, tenant_id: str, user_id: str, timeout: float) -> None:
        self.base_url = base_url.rstrip("/")
        self.session = requests.Session()
        self.session.headers.update({"X-Tenant-ID": tenant_id, "X-User-ID": user_id})
        self.timeout = timeout
        self.max_rate_limit_retries = 5

    @staticmethod
    def _retry_delay(resp: Any, attempt: int) -> float:
        retry_after = str(resp.headers.get("Retry-After") or "").strip()
        if retry_after:
            try:
                return max(0.1, min(5.0, float(retry_after)))
            except ValueError:
                pass
        try:
            payload = resp.json()
            detail = payload.get("detail") if isinstance(payload, dict) else {}
            if isinstance(detail, dict):
                return max(0.1, min(5.0, float(detail.get("retry_after_sec") or 0)))
        except Exception:
            pass
        return min(5.0, 0.5 * (2 ** max(0, int(attempt))))

    def request(self, method: str, path: str, **kwargs: Any) -> tuple[Any, float]:
        url = f"{self.base_url}{path}"
        started = time.perf_counter()
        resp: Any = None
        for attempt in range(self.max_rate_limit_retries + 1):
            resp = self.session.request(method, url, timeout=self.timeout, **kwargs)
            if resp.status_code != 429 or attempt >= self.max_rate_limit_retries:
                break
            time.sleep(self._retry_delay(resp, attempt))
        elapsed_ms = (time.perf_counter() - started) * 1000.0
        if resp is None:  # defensive; loop always executes at least once.
            fail(f"{method} {path} did not return a response")
        return resp, elapsed_ms

    def json(
        self,
        method: str,
        path: str,
        expected: set[int] | None = None,
        **kwargs: Any,
    ) -> tuple[dict[str, Any], float]:
        resp, elapsed_ms = self.request(method, path, **kwargs)
        expected = expected or {200}
        if resp.status_code not in expected:
            fail(f"{method} {path} returned {resp.status_code}: {resp.text[:800]}")
        try:
            data = resp.json()
        except Exception as exc:  # noqa: BLE001
            fail(f"{method} {path} returned non-JSON: {exc}; body={resp.text[:800]}")
        if not isinstance(data, dict):
            fail(f"{method} {path} returned non-object JSON: {type(data).__name__}")
        return data, elapsed_ms


def download_corpus(corpus_dir: Path, evidence: Evidence) -> list[Path]:
    corpus_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    for source in SOURCES:
        out = corpus_dir / source.filename
        if not out.exists() or out.stat().st_size < 64:
            response = requests.get(
                source.url,
                timeout=60,
                headers={"User-Agent": "MimirQ-production-readiness/1.0"},
            )
            if response.status_code != 200:
                fail(f"download failed {source.url}: HTTP {response.status_code}")
            out.write_bytes(response.content)
        size = out.stat().st_size
        if size < 256:
            fail(f"downloaded file is too small: {out} ({size} bytes)")
        evidence.downloads.append(
            {
                "filename": source.filename,
                "kind": source.kind,
                "url": source.url,
                "bytes": size,
            }
        )
        paths.append(out)
    return paths


def read_text_sample(path: Path, max_chars: int = 12_000) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw.decode(encoding, errors="ignore")[:max_chars]
        except Exception:
            continue
    return raw.decode("utf-8", errors="ignore")[:max_chars]


def generate_office_files(corpus_dir: Path, evidence: Evidence) -> list[Path]:
    rfc9000 = corpus_dir / "rfc9000-quic.txt"
    fastapi = corpus_dir / "fastapi-readme.md"
    iris = corpus_dir / "iris.csv"
    docx_path = corpus_dir / "mixed-rag-operations-brief.docx"
    xlsx_path = corpus_dir / "iris-quality-sample.xlsx"

    doc = Document()
    doc.add_heading("MimirQ Production Readiness Brief", level=1)
    doc.add_paragraph("This DOCX is generated from real downloaded public documents to validate Office parsing.")
    doc.add_heading("QUIC sample from RFC 9000", level=2)
    doc.add_paragraph(read_text_sample(rfc9000, 4500))
    doc.add_heading("FastAPI README sample", level=2)
    doc.add_paragraph(read_text_sample(fastapi, 4500))
    doc.add_paragraph(
        "Business check: retrieval should find QUIC, FastAPI and operational parsing signals in this file."
    )
    doc.save(docx_path)

    wb = Workbook()
    ws = wb.active
    ws.title = "iris"
    rows = list(csv.reader(iris.read_text(encoding="utf-8", errors="ignore").splitlines()))
    for row in rows[:151]:
        ws.append(row)
    summary = wb.create_sheet("quality_notes")
    summary.append(["metric", "value"])
    summary.append(["source_url", "https://raw.githubusercontent.com/mwaskom/seaborn-data/master/iris.csv"])
    summary.append(["rows_imported", max(0, len(rows) - 1)])
    summary.append(
        [
            "purpose",
            "validate spreadsheet parsing, chunking and retrieval over structured rows",
        ]
    )
    wb.save(xlsx_path)

    out = [docx_path, xlsx_path]
    for path in out:
        evidence.generated_files.append({"filename": path.name, "bytes": path.stat().st_size})
    return out


def ensure_runtime_settings(api: Api, evidence: Evidence) -> None:
    current, _elapsed = api.json("GET", "/api/v1/settings")
    feature_flags = dict(current.get("feature_flags") or {})
    kg = dict(current.get("kg") or {})
    changed = False
    if feature_flags.get("kg_enabled") is not True:
        feature_flags["kg_enabled"] = True
        changed = True
    if kg.get("chat_enabled") is not True:
        kg["chat_enabled"] = True
        changed = True
    if changed:
        payload = {"feature_flags": feature_flags, "kg": kg}
        updated, elapsed_ms = api.json("PUT", "/api/v1/settings", json=payload)
        evidence.settings["update"] = {"elapsed_ms": round(elapsed_ms, 2), "response": updated}
        time.sleep(0.2)
    verified, _ = api.json("GET", "/api/v1/settings")
    evidence.settings["effective"] = {
        "kg_enabled": (verified.get("feature_flags") or {}).get("kg_enabled"),
        "kg_chat_enabled": (verified.get("kg") or {}).get("chat_enabled"),
        "embedding_provider": (verified.get("embedding") or {}).get("provider"),
        "embedding_model": (verified.get("embedding") or {}).get("model"),
        "llm_model": (verified.get("llm") or {}).get("model"),
        "vector_backend": (verified.get("milvus") or {}).get("host"),
        "default_parser_backend": (verified.get("rag") or {}).get("default_parser_backend"),
        "default_chunk_strategy": (verified.get("rag") or {}).get("default_chunk_strategy"),
    }
    evidence.check(
        "settings_kg_enabled",
        (
            evidence.settings["effective"]["kg_enabled"] is True
            and evidence.settings["effective"]["kg_chat_enabled"] is True
        ),
        effective=evidence.settings["effective"],
    )


def probe_llm_provider(api: Api, evidence: Evidence, *, timeout_sec: float = 15.0) -> None:
    """Record whether the configured LLM provider is genuinely callable.

    The settings endpoint returns masked secrets, so the runner uses the local
    process/.env value for the probe and only writes redacted metadata.
    """
    current, _ = api.json("GET", "/api/v1/settings")
    llm = dict(current.get("llm") or {})
    api_base = str(llm.get("api_base") or load_local_env_value("LLM_API_BASE")).strip()
    model = str(llm.get("model") or load_local_env_value("LLM_MODEL")).strip()
    api_key, api_key_source = load_llm_probe_api_key(api_base)
    if not api_key or not api_base or not model:
        evidence.provider_health = {
            "checked": False,
            "success": False,
            "reason": "missing_local_llm_api_key_api_base_or_model",
            "api_key_source": api_key_source,
            "model": model or None,
            "api_base_configured": bool(api_base),
        }
        evidence.check("llm_provider_health_recorded", True, provider_health=evidence.provider_health)
        return

    payload = {
        "api_key": api_key,
        "api_base": api_base,
        "model": model,
        "temperature": 0,
        "timeout": max(1.0, float(timeout_sec)),
        "max_retries": 0,
    }
    response, elapsed_ms = api.request("POST", "/api/v1/settings/llm/test", json=payload)
    try:
        data = response.json()
    except Exception:
        data = {"success": False, "message": response.text[:400]}

    success = response.status_code == 200 and bool(data.get("success"))
    evidence.provider_health = {
        "checked": True,
        "success": success,
        "status_code": response.status_code,
        "elapsed_ms": round(elapsed_ms, 2),
        "model": model,
        "api_base": api_base,
        "api_key_source": api_key_source,
        "message": str(data.get("message") or "")[:400],
    }
    evidence.check("llm_provider_health_recorded", True, provider_health=evidence.provider_health)


def create_dataset(api: Api, evidence: Evidence) -> str:
    payload = {
        "name": f"Prod Readiness Corpus {now_id()}",
        "description": "Automated 10+ document production-readiness chain: parsing, chunking, KG, RAG and chat.",
        "default_parser_backend": "auto",
        "default_chunk_strategy": "langchain_recursive",
        "pipeline": {
            "governance_enabled": True,
            "governance_remove_noise_lines": True,
            "governance_unwrap_lines": True,
            "governance_drop_duplicate_paragraphs": True,
            "persist_parsed_content": True,
            "persist_parsed_content_max_chars": 500000,
            "chunk_size": 1600,
            "chunk_overlap": 160,
            "chunk_vector_enabled": True,
            "bm25_index_enabled": True,
            "kg_enabled": False,
            "event_vector_enabled": False,
            "entity_vector_enabled": False,
        },
        "rag_defaults": {
            "top_k": 6,
            "score_threshold": 0.0,
            "retrieval_mode": "hybrid",
            "enable_reranker": False,
            "enable_multi_query": False,
            "enable_hyde": False,
            "enable_query_decomposition": False,
        },
    }
    data, elapsed_ms = api.json("POST", "/api/v1/datasets/", expected={201}, json=payload)
    dataset_id = str(data.get("id") or "")
    if not dataset_id:
        fail("dataset create response missing id")
    evidence.dataset = {"elapsed_ms": round(elapsed_ms, 2), **data}
    evidence.check("dataset_created", True, dataset_id=dataset_id)
    return dataset_id


def wait_for_document_terminal(
    api: Api,
    doc_id: str,
    *,
    timeout_sec: float,
    poll_interval_sec: float = 2.0,
) -> dict[str, Any]:
    deadline = time.monotonic() + timeout_sec
    last: dict[str, Any] = {}
    while time.monotonic() < deadline:
        data, elapsed_ms = api.json("GET", f"/api/v1/documents/{doc_id}")
        last = data
        last["_detail_elapsed_ms"] = round(elapsed_ms, 2)
        if str(data.get("status") or "") in {"completed", "failed", "quarantined", "cancelled"}:
            return last
        time.sleep(max(0.1, float(poll_interval_sec)))
    fail(f"document processing timed out for {doc_id}: last_status={last.get('status') if last else 'unknown'}")


def upload_documents(
    api: Api,
    dataset_id: str,
    files: list[Path],
    evidence: Evidence,
    *,
    per_upload_timeout_sec: float = 0.0,
) -> list[str]:
    doc_ids: list[str] = []
    for path in files:
        data = {
            "dataset_id": dataset_id,
            "parser_backend": "auto",
            "chunk_strategy": "langchain_recursive",
            "governance_enabled": "true",
            "chunk_vector_enabled": "true",
            "bm25_index_enabled": "true",
            "kg_enabled": "false",
            "event_vector_enabled": "false",
            "entity_vector_enabled": "false",
            "user_metadata": json.dumps(
                {
                    "production_readiness": True,
                    "source_file": path.name,
                    "corpus": "production-readiness-chain",
                },
                ensure_ascii=False,
            ),
        }
        with path.open("rb") as fh:
            response, elapsed_ms = api.request(
                "POST",
                "/api/v1/documents/upload",
                data=data,
                files={"file": (path.name, fh)},
            )
        if response.status_code not in {200, 201}:
            fail(f"upload failed for {path.name}: HTTP {response.status_code} {response.text[:800]}")
        payload = response.json()
        doc_id = str(payload.get("id") or "")
        if not doc_id:
            fail(f"upload response missing id for {path.name}")
        doc_ids.append(doc_id)
        upload_row: dict[str, Any] = {
            "filename": path.name,
            "document_id": doc_id,
            "status": payload.get("status"),
            "file_type": payload.get("file_type"),
            "file_size": payload.get("file_size"),
            "elapsed_ms": round(elapsed_ms, 2),
            "metadata": payload.get("doc_metadata") or payload.get("metadata") or {},
        }
        if per_upload_timeout_sec > 0:
            # Keep the readiness chain deterministic on local/dev stacks where
            # background parsing + KG extraction can otherwise exhaust DB pool
            # capacity before later upload requests even resolve their dataset.
            terminal = wait_for_document_terminal(api, doc_id, timeout_sec=per_upload_timeout_sec)
            upload_row["terminal_status"] = terminal.get("status")
            upload_row["terminal_wait_ms"] = terminal.get("_detail_elapsed_ms")
        evidence.uploads.append(upload_row)
    evidence.check("uploaded_at_least_10_documents", len(doc_ids) >= 10, count=len(doc_ids))
    return doc_ids


def wait_for_documents(api: Api, doc_ids: list[str], evidence: Evidence, timeout_sec: float) -> None:
    deadline = time.monotonic() + timeout_sec
    pending = set(doc_ids)
    final: dict[str, dict[str, Any]] = {}
    while pending and time.monotonic() < deadline:
        for doc_id in list(pending):
            data, elapsed_ms = api.json("GET", f"/api/v1/documents/{doc_id}")
            status = str(data.get("status") or "")
            if status in {"completed", "failed", "quarantined", "cancelled"}:
                data["_detail_elapsed_ms"] = round(elapsed_ms, 2)
                final[doc_id] = data
                pending.remove(doc_id)
        if pending:
            time.sleep(2)
    if pending:
        fail(f"document processing timed out for {len(pending)} docs: {sorted(pending)[:5]}")

    for doc_id in doc_ids:
        doc = final[doc_id]
        status = str(doc.get("status") or "")
        chunks, chunk_elapsed_ms = api.json("GET", f"/api/v1/documents/{doc_id}/chunks?limit=2000")
        parsed, parsed_elapsed_ms = api.json("GET", f"/api/v1/documents/{doc_id}/parsed-content?max_chars=20000")
        chunk_total = int(chunks.get("total") or 0)
        evidence.documents.append(
            {
                "document_id": doc_id,
                "filename": doc.get("filename"),
                "file_type": doc.get("file_type"),
                "file_size": doc.get("file_size"),
                "status": status,
                "processing_progress": doc.get("processing_progress"),
                "chunk_total": chunk_total,
                "parsed_content_available": bool(parsed.get("available")),
                "parsed_markdown_chars": len(str(parsed.get("markdown_content") or "")),
                "metadata": doc.get("doc_metadata") or doc.get("metadata") or {},
                "timings_ms": {
                    "detail": doc.get("_detail_elapsed_ms"),
                    "chunks": round(chunk_elapsed_ms, 2),
                    "parsed_content": round(parsed_elapsed_ms, 2),
                },
            }
        )

    failed = [d for d in evidence.documents if d["status"] != "completed"]
    zero_chunks = [d for d in evidence.documents if int(d["chunk_total"] or 0) <= 0]
    evidence.check("all_documents_completed", not failed, failures=failed[:5])
    evidence.check("all_documents_have_chunks", not zero_chunks, failures=zero_chunks[:5])


def run_chunk_previews(api: Api, dataset_id: str, corpus_files: list[Path], evidence: Evidence) -> None:
    candidates = [
        ("langchain_recursive", corpus_files[0]),
        ("markdown", next(p for p in corpus_files if p.suffix == ".md")),
        ("semantic_sentence", corpus_files[1]),
        ("parent_child", corpus_files[0]),
        ("csv_rows", next(p for p in corpus_files if p.suffix == ".csv")),
    ]
    for strategy, path in candidates:
        data = {
            "dataset_id": dataset_id,
            "parser_backend": "auto",
            "chunk_strategy": strategy,
            "chunk_size": "1200",
            "chunk_overlap": "120",
            "include_original_text": "false",
            "include_chunks": "true",
            "max_chunks": "50",
        }
        attempts: list[dict[str, Any]] = []
        payload: dict[str, Any] = {}
        response_text = ""
        response_status = 0
        elapsed_ms = 0.0
        count = 0
        for attempt in range(1, 4):
            with path.open("rb") as fh:
                response, elapsed_ms = api.request(
                    "POST",
                    "/api/v1/documents/chunk-preview",
                    data=data,
                    files={"file": (path.name, fh)},
                )
            response_status = int(response.status_code)
            response_text = response.text[:800]
            payload = response.json() if response.status_code == 200 else {}
            chunks = payload.get("chunks") if isinstance(payload, dict) else None
            count = len(chunks or [])
            attempts.append(
                {
                    "attempt": attempt,
                    "status_code": response_status,
                    "chunk_count": count,
                    "elapsed_ms": round(elapsed_ms, 2),
                }
            )
            if response_status == 200 and count > 0:
                break
            if attempt < 3 and (
                response_status in {429, 500, 502, 503, 504} or (response_status == 200 and count == 0)
            ):
                time.sleep(0.5 * attempt)
                continue
            break
        if response_status != 200:
            evidence.chunk_previews.append(
                {
                    "strategy": strategy,
                    "filename": path.name,
                    "ok": False,
                    "status_code": response_status,
                    "body": response_text,
                    "elapsed_ms": round(elapsed_ms, 2),
                    "attempts": attempts,
                }
            )
            continue
        evidence.chunk_previews.append(
            {
                "strategy": strategy,
                "filename": path.name,
                "ok": count > 0,
                "chunk_count": count,
                "elapsed_ms": round(elapsed_ms, 2),
                "metrics": payload.get("metrics") or payload.get("stats") or {},
                "attempts": attempts,
            }
        )
    bad = [row for row in evidence.chunk_previews if not row.get("ok")]
    evidence.check("chunk_preview_strategies_work", not bad, failures=bad)


def _kg_stats_path(dataset_id: str) -> str:
    return f"/api/v1/kg/stats?{urlencode({'dataset_id': dataset_id})}"


def _kg_extract_path(doc_id: str) -> str:
    params = {
        "replace_existing": "true",
        "extract_relations": "false",
        "extract_skills": "false",
        "extraction_backend": "heuristic",
    }
    return f"/api/v1/kg/documents/{doc_id}/extract?{urlencode(params)}"


def _kg_result_count(row: dict[str, Any]) -> int:
    return int(row.get("events") or 0) + int(row.get("entities") or 0)


def _kg_search_ready(row: dict[str, Any]) -> bool:
    return float(row.get("elapsed_ms") or 0) <= 3000.0 and _kg_result_count(row) > 0


def _kg_search(api: Api, dataset_id: str, query: str) -> dict[str, Any]:
    payload = {"query": query, "dataset_id": dataset_id}
    result, kg_elapsed_ms = api.json("POST", "/api/v1/kg/search", json=payload)
    raw = result.get("result") or {}
    return {
        "query": query,
        "elapsed_ms": round(kg_elapsed_ms, 2),
        "events": len(raw.get("events") or []),
        "entities": len(raw.get("entities") or []),
        "stats": raw.get("stats") or {},
    }


def _kg_search_with_retry(api: Api, dataset_id: str, query: str) -> dict[str, Any]:
    attempts: list[dict[str, Any]] = []
    final_row: dict[str, Any] = {}
    for attempt in range(1, 3):
        row = _kg_search(api, dataset_id, query)
        attempts.append(
            {
                "attempt": attempt,
                "elapsed_ms": row["elapsed_ms"],
                "events": row["events"],
                "entities": row["entities"],
            }
        )
        final_row = row
        if _kg_search_ready(row):
            break
        if attempt < 2:
            time.sleep(0.25)
    final_row["attempts"] = attempts
    return final_row


def _kg_warmup_rows(api: Api, dataset_id: str) -> list[dict[str, Any]]:
    warmup_rows: list[dict[str, Any]] = []
    query = "Warm up KG search over QUIC FastAPI accessibility"
    for attempt in range(1, 4):
        row = _kg_search(api, dataset_id, query)
        warmup_rows.append({"attempt": attempt, **row})
        if _kg_search_ready(row):
            break
    return warmup_rows


def _manual_kg_extracts(api: Api, doc_ids: list[str]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for doc_id in doc_ids:
        try:
            extracted, extract_elapsed_ms = api.json(
                "POST",
                _kg_extract_path(doc_id),
                expected={200},
            )
            rows.append(
                {
                    "document_id": doc_id,
                    "elapsed_ms": round(extract_elapsed_ms, 2),
                    **extracted,
                }
            )
        except Exception as exc:  # noqa: BLE001
            rows.append({"document_id": doc_id, "error": str(exc)})
    return rows


def ensure_kg(api: Api, dataset_id: str, doc_ids: list[str], evidence: Evidence) -> None:
    stats, elapsed_ms = api.json("GET", _kg_stats_path(dataset_id))
    manual_extracts: list[dict[str, Any]] = []
    if int(stats.get("events") or 0) <= 0 or int(stats.get("entities") or 0) <= 0:
        manual_extracts = _manual_kg_extracts(api, doc_ids)
        stats, elapsed_ms = api.json("GET", _kg_stats_path(dataset_id))

    warmup_rows = _kg_warmup_rows(api, dataset_id)
    search_queries = ["QUIC transport handshake", "FastAPI HTTP client", "accessibility conformance"]
    search_rows = [_kg_search_with_retry(api, dataset_id, query) for query in search_queries]

    evidence.kg = {
        "stats_elapsed_ms": round(elapsed_ms, 2),
        "stats": stats,
        "manual_extract_backend": "heuristic",
        "manual_extracts": manual_extracts,
        "warmup": warmup_rows,
        "search": search_rows,
    }
    evidence.check(
        "kg_has_entities_and_events",
        int(stats.get("events") or 0) > 0 and int(stats.get("entities") or 0) > 0,
        stats=stats,
    )
    latest_warmup = warmup_rows[-1] if warmup_rows else {}
    evidence.check(
        "kg_search_warmup_completed",
        bool(latest_warmup) and _kg_search_ready(latest_warmup),
        latest=latest_warmup,
    )
    slow = [row for row in search_rows if float(row.get("elapsed_ms") or 0) > 3000.0]
    empty = [row for row in search_rows if _kg_result_count(row) <= 0]
    evidence.check("kg_search_under_3s", not slow, slow=slow)
    evidence.check("kg_search_returns_results", not empty, empty=empty)


def run_retrieval(api: Api, dataset_id: str, evidence: Evidence) -> None:
    queries = [
        "What does RFC 9000 say QUIC provides compared with TCP?",
        "How does FastAPI describe API documentation and validation?",
        "What does WCAG say about accessibility conformance?",
        "What species are present in the iris sample data?",
    ]
    for query in queries:
        payload = {
            "query": query,
            "dataset_id": dataset_id,
            "rag_config": {
                "top_k": 6,
                "score_threshold": 0.0,
                "retrieval_mode": "hybrid",
                "enable_reranker": False,
                "enable_multi_query": False,
                "enable_hyde": False,
                "enable_query_decomposition": False,
            },
        }
        attempts: list[dict[str, Any]] = []
        final_row: dict[str, Any] = {}
        for attempt in range(1, 3):
            result, elapsed_ms = api.json("POST", "/api/v1/rag/retrieve-preview", json=payload)
            citations = result.get("citations") or []
            row = {
                "query": query,
                "elapsed_ms": round(elapsed_ms, 2),
                "citation_count": len(citations),
                "metrics": result.get("metrics") or {},
                "top_sources": [
                    {
                        "document_id": (c.get("document_id") or c.get("metadata", {}).get("document_id")),
                        "source": c.get("source") or c.get("filename") or c.get("metadata", {}).get("source"),
                        "score": c.get("score"),
                    }
                    for c in citations[:5]
                    if isinstance(c, dict)
                ],
            }
            attempts.append(
                {
                    "attempt": attempt,
                    "elapsed_ms": row["elapsed_ms"],
                    "citation_count": row["citation_count"],
                }
            )
            final_row = row
            if elapsed_ms <= 3000.0 and len(citations) > 0:
                break
            if attempt < 2:
                time.sleep(0.25)
        final_row["attempts"] = attempts
        evidence.retrieval.append(final_row)
    slow = [row for row in evidence.retrieval if float(row.get("elapsed_ms") or 0) > 3000.0]
    empty = [row for row in evidence.retrieval if int(row.get("citation_count") or 0) <= 0]
    evidence.check("rag_retrieval_under_3s", not slow, slow=slow)
    evidence.check("rag_retrieval_returns_citations", not empty, empty=empty)


def warm_retrieval_path(api: Api, dataset_id: str, evidence: Evidence) -> None:
    """Warm retrieval indexes/providers before measuring the production SLO path."""
    query = "Warm up QUIC FastAPI retrieval path"
    payload = {
        "query": query,
        "dataset_id": dataset_id,
        "rag_config": {
            "top_k": 6,
            "score_threshold": 0.0,
            "retrieval_mode": "hybrid",
            "enable_reranker": False,
            "enable_multi_query": False,
            "enable_hyde": False,
            "enable_query_decomposition": False,
        },
    }
    for attempt in range(1, 4):
        result, elapsed_ms = api.json("POST", "/api/v1/rag/retrieve-preview", json=payload)
        row = {
            "attempt": attempt,
            "elapsed_ms": round(elapsed_ms, 2),
            "citation_count": len(result.get("citations") or []),
            "metrics": result.get("metrics") or {},
        }
        evidence.retrieval_warmups.append(row)
        if elapsed_ms <= 3000.0 and row["citation_count"] > 0:
            break
    latest = evidence.retrieval_warmups[-1] if evidence.retrieval_warmups else {}
    evidence.check(
        "rag_retrieval_warmup_completed",
        bool(latest) and float(latest.get("elapsed_ms") or 0) <= 3000.0 and int(latest.get("citation_count") or 0) > 0,
        latest=latest,
    )


def run_chat(api: Api, dataset_id: str, evidence: Evidence) -> None:
    questions = [
        "请用三点概括这个数据集中 QUIC、HTTP 语义和 FastAPI 文档的共同技术主题。",
        "基于当前知识库，说明 iris 数据集中包含哪些物种，并指出它属于什么类型的数据。",
    ]
    for question in questions:
        payload = {
            "message": question,
            "dataset_id": dataset_id,
            "stream": False,
            "rag_config": {
                "top_k": 6,
                "score_threshold": 0.0,
                "retrieval_mode": "hybrid",
                "use_graph": True,
                "enable_reranker": False,
                "enable_multi_query": False,
                "enable_hyde": False,
                "enable_query_decomposition": False,
                "max_tokens": 1200,
                "answer_mode": "extractive",
            },
        }
        result, elapsed_ms = api.json("POST", "/api/v1/chat", json=payload)
        content = str(result.get("content") or "")
        citations = result.get("citations") or []
        evidence.chat.append(
            {
                "question": question,
                "elapsed_ms": round(elapsed_ms, 2),
                "content_chars": len(content),
                "citation_count": len(citations),
                "fallback_used": bool((result.get("metrics") or {}).get("generation_fallback_used")),
                "fallback_reason": (result.get("metrics") or {}).get("generation_fallback_reason"),
                "retrieval_mode": result.get("retrieval_mode"),
                "vector_backend": result.get("vector_backend"),
                "metrics": result.get("metrics") or {},
                "answer_excerpt": content[:500],
            }
        )
    bad = [
        row
        for row in evidence.chat
        if int(row.get("content_chars") or 0) <= 20 or int(row.get("citation_count") or 0) <= 0
    ]
    evidence.check("chat_answers_with_citations", not bad, failures=bad)


def run_default_chat_degradation(api: Api, dataset_id: str, evidence: Evidence) -> None:
    """Verify default LLM mode answers with citations, or degrades fast when provider is unavailable."""
    questions = [
        "默认问答模式下，这批 QUIC 和 FastAPI 资料可以验证哪些 RAG 能力？",
        "默认问答模式下，iris 文件说明了什么数据类型？",
    ]
    for question in questions:
        payload = {
            "message": question,
            "dataset_id": dataset_id,
            "stream": False,
            "rag_config": {
                "top_k": 6,
                "score_threshold": 0.0,
                "retrieval_mode": "hybrid",
                "use_graph": True,
                "enable_reranker": False,
                "enable_multi_query": False,
                "enable_hyde": False,
                "enable_query_decomposition": False,
                "max_tokens": 1200,
            },
        }
        result, elapsed_ms = api.json("POST", "/api/v1/chat", json=payload)
        content = str(result.get("content") or "")
        citations = result.get("citations") or []
        metrics = result.get("metrics") or {}
        evidence.default_chat.append(
            {
                "question": question,
                "elapsed_ms": round(elapsed_ms, 2),
                "content_chars": len(content),
                "citation_count": len(citations),
                "fallback_used": bool(metrics.get("generation_fallback_used")),
                "fallback_reason": metrics.get("generation_fallback_reason"),
                "fallback_error": str(metrics.get("generation_fallback_error") or "")[:400],
                "metrics": metrics,
                "answer_excerpt": content[:500],
            }
        )
    provider_ok = bool(evidence.provider_health.get("success"))
    bad: list[dict[str, Any]] = []
    for row in evidence.default_chat:
        has_answer = int(row.get("content_chars") or 0) > 20
        has_citations = int(row.get("citation_count") or 0) > 0
        if not has_answer or not has_citations:
            bad.append({**row, "reason": "missing_answer_or_citations"})
            continue
        if provider_ok:
            continue
        if not bool(row.get("fallback_used")):
            bad.append({**row, "reason": "provider_unavailable_without_fallback"})
            continue
        if float(row.get("elapsed_ms") or 0) > 3000.0:
            bad.append({**row, "reason": "fallback_too_slow"})
    evidence.check(
        "default_chat_answers_or_degrades_with_citations",
        not bad,
        provider_ok=provider_ok,
        failures=bad,
    )


def summarize_generation_readiness(evidence: Evidence) -> None:
    """Gate production-readiness on either live LLM or verified fast degradation."""
    provider_ok = bool(evidence.provider_health.get("success"))
    degraded_ok = bool(
        evidence.default_chat
        and all(
            int(row.get("content_chars") or 0) > 20
            and int(row.get("citation_count") or 0) > 0
            and bool(row.get("fallback_used"))
            and float(row.get("elapsed_ms") or 0) <= 3000.0
            for row in evidence.default_chat
        )
    )
    evidence.check(
        "llm_generation_available_or_fast_degraded",
        provider_ok or degraded_ok,
        provider_health=evidence.provider_health,
        default_chat=[
            {
                "elapsed_ms": row.get("elapsed_ms"),
                "citation_count": row.get("citation_count"),
                "fallback_used": row.get("fallback_used"),
                "fallback_reason": row.get("fallback_reason"),
                "fallback_error": row.get("fallback_error"),
            }
            for row in evidence.default_chat
        ],
    )


def summarize_formats(evidence: Evidence) -> None:
    formats = sorted({str(row.get("file_type") or "").lower() for row in evidence.documents if row.get("file_type")})
    evidence.check("multiple_file_types_detected", len(formats) >= 5, formats=formats, count=len(formats))
    total_chunks = sum(int(row.get("chunk_total") or 0) for row in evidence.documents)
    evidence.check("nontrivial_chunk_volume", total_chunks >= 20, total_chunks=total_chunks)


def summarize_runtime_quality(evidence: Evidence) -> None:
    expansion_leaks: list[dict[str, Any]] = []
    for scope, rows in (
        ("retrieval_warmup", evidence.retrieval_warmups),
        ("retrieval", evidence.retrieval),
        ("chat", evidence.chat),
        ("default_chat", evidence.default_chat),
    ):
        for index, row in enumerate(rows, start=1):
            metrics = row.get("metrics") or {}
            if not isinstance(metrics, dict):
                continue
            if (
                bool(metrics.get("multi_query_enabled"))
                or bool(metrics.get("hyde_enabled"))
                or bool(metrics.get("multi_query_parse_error"))
                or bool(metrics.get("hyde_parse_error"))
            ):
                expansion_leaks.append(
                    {
                        "scope": scope,
                        "index": index,
                        "multi_query_enabled": metrics.get("multi_query_enabled"),
                        "hyde_enabled": metrics.get("hyde_enabled"),
                        "multi_query_parse_error": metrics.get("multi_query_parse_error"),
                        "hyde_parse_error": metrics.get("hyde_parse_error"),
                    }
                )
    evidence.check("retrieval_path_has_no_llm_expansion_errors", not expansion_leaks, leaks=expansion_leaks)


def write_report(evidence: Evidence) -> None:
    out_dir = Path(evidence.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    report_json = out_dir / "report.json"
    report_md = out_dir / "report.md"
    report_json.write_text(json_dumps(evidence.__dict__), encoding="utf-8")

    lines = [
        "# MimirQ Production Readiness Chain Evidence",
        "",
        f"- Started: `{evidence.started_at}`",
        f"- Base URL: `{evidence.base_url}`",
        f"- Dataset: `{evidence.dataset.get('name')}` / `{evidence.dataset.get('id')}`",
        f"- Corpus dir: `{evidence.corpus_dir}`",
        "",
        "## Gate Summary",
        "",
    ]
    for check in evidence.checks:
        details = {key: value for key, value in check.items() if key not in {"name", "ok"}}
        marker = "x" if check.get("ok") else " "
        lines.append(f"- [{marker}] {check.get('name')} `{json.dumps(details, ensure_ascii=False, default=str)}`")
    lines.extend(
        [
            "",
            "## Documents",
            "",
            "| file | type | status | chunks | parsed chars |",
            "| --- | --- | --- | ---: | ---: |",
        ]
    )
    for doc in evidence.documents:
        lines.append(
            f"| {doc.get('filename')} | {doc.get('file_type')} "
            f"| {doc.get('status')} | {doc.get('chunk_total')} "
            f"| {doc.get('parsed_markdown_chars')} |"
        )
    lines.extend(
        [
            "",
            "## Retrieval Latency",
            "",
            "| query | elapsed ms | citations |",
            "| --- | ---: | ---: |",
        ]
    )
    for row in evidence.retrieval:
        lines.append(f"| {row.get('query')} | {row.get('elapsed_ms')} | {row.get('citation_count')} |")
    lines.extend(
        [
            "",
            "## Retrieval Warmup",
            "",
            "| attempt | elapsed ms | citations |",
            "| ---: | ---: | ---: |",
        ]
    )
    for row in evidence.retrieval_warmups:
        lines.append(f"| {row.get('attempt')} | {row.get('elapsed_ms')} | {row.get('citation_count')} |")
    lines.extend(
        [
            "",
            "## KG Search Warmup",
            "",
            "| attempt | elapsed ms | events | entities |",
            "| ---: | ---: | ---: | ---: |",
        ]
    )
    for row in evidence.kg.get("warmup", []):
        lines.append(
            f"| {row.get('attempt')} | {row.get('elapsed_ms')} | {row.get('events')} | {row.get('entities')} |"
        )
    lines.extend(
        [
            "",
            "## KG Search Latency",
            "",
            "| query | elapsed ms | events | entities |",
            "| --- | ---: | ---: | ---: |",
        ]
    )
    for row in evidence.kg.get("search", []):
        lines.append(f"| {row.get('query')} | {row.get('elapsed_ms')} | {row.get('events')} | {row.get('entities')} |")
    lines.extend(
        [
            "",
            "## LLM Provider Health",
            "",
            f"- Checked: `{evidence.provider_health.get('checked')}`",
            f"- Success: `{evidence.provider_health.get('success')}`",
            f"- Model: `{evidence.provider_health.get('model')}`",
            f"- Elapsed ms: `{evidence.provider_health.get('elapsed_ms')}`",
            f"- Message: `{evidence.provider_health.get('message') or evidence.provider_health.get('reason') or ''}`",
        ]
    )
    lines.extend(
        [
            "",
            "## Chat",
            "",
            "| question | elapsed ms | chars | citations |",
            "| --- | ---: | ---: | ---: |",
        ]
    )
    for row in evidence.chat:
        lines.append(
            f"| {row.get('question')} | {row.get('elapsed_ms')} "
            f"| {row.get('content_chars')} | {row.get('citation_count')} |"
        )
    lines.extend(
        [
            "",
            "## Default Chat",
            "",
            "| question | elapsed ms | chars | citations | reason |",
            "| --- | ---: | ---: | ---: | --- |",
        ]
    )
    for row in evidence.default_chat:
        lines.append(
            f"| {row.get('question')} | {row.get('elapsed_ms')} "
            f"| {row.get('content_chars')} | {row.get('citation_count')} "
            f"| {row.get('fallback_reason')} |"
        )
    if evidence.failures:
        lines.extend(["", "## Failures", ""])
        lines.extend([f"- {failure}" for failure in evidence.failures])
    report_md.write_text("\n".join(lines) + "\n", encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--base-url", default="http://127.0.0.1:8000")
    parser.add_argument("--tenant-id", default=TENANT_ID)
    parser.add_argument("--user-id", default=USER_ID)
    parser.add_argument("--timeout", type=float, default=180.0)
    parser.add_argument("--processing-timeout", type=float, default=1800.0)
    parser.add_argument(
        "--per-upload-timeout",
        type=float,
        default=None,
        help=(
            "Optional per-upload terminal wait. Defaults to "
            "processing-timeout; set 0 to defer waiting until the "
            "bulk wait phase."
        ),
    )
    parser.add_argument(
        "--llm-probe-timeout",
        type=float,
        default=15.0,
        help="Configured provider call timeout for /settings/llm/test.",
    )
    parser.add_argument("--corpus-dir", default="")
    parser.add_argument("--output-dir", default="")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    run_id = now_id()
    root = Path("artifacts/production-readiness") / run_id
    corpus_dir = Path(args.corpus_dir) if args.corpus_dir else root / "corpus"
    output_dir = Path(args.output_dir) if args.output_dir else root
    evidence = Evidence(
        started_at=datetime.now(UTC).isoformat(),
        base_url=args.base_url.rstrip("/"),
        tenant_id=args.tenant_id,
        user_id=args.user_id,
        corpus_dir=str(corpus_dir.resolve()),
        output_dir=str(output_dir.resolve()),
    )

    try:
        downloaded = download_corpus(corpus_dir, evidence)
        generated = generate_office_files(corpus_dir, evidence)
        files = downloaded + generated
        api = Api(args.base_url, args.tenant_id, args.user_id, timeout=float(args.timeout))
        ensure_runtime_settings(api, evidence)
        probe_llm_provider(api, evidence, timeout_sec=float(args.llm_probe_timeout))
        dataset_id = create_dataset(api, evidence)
        per_upload_timeout = (
            float(args.processing_timeout) if args.per_upload_timeout is None else float(args.per_upload_timeout)
        )
        doc_ids = upload_documents(
            api,
            dataset_id,
            files,
            evidence,
            per_upload_timeout_sec=per_upload_timeout,
        )
        wait_for_documents(api, doc_ids, evidence, timeout_sec=float(args.processing_timeout))
        run_chunk_previews(api, dataset_id, files, evidence)
        ensure_kg(api, dataset_id, doc_ids, evidence)
        warm_retrieval_path(api, dataset_id, evidence)
        run_retrieval(api, dataset_id, evidence)
        run_chat(api, dataset_id, evidence)
        run_default_chat_degradation(api, dataset_id, evidence)
        summarize_generation_readiness(evidence)
        summarize_formats(evidence)
        summarize_runtime_quality(evidence)
    except Exception as exc:  # noqa: BLE001
        evidence.failures.append(str(exc))
        print(f"ERROR: {exc}", file=sys.stderr)
    finally:
        write_report(evidence)

    print(json_dumps({"output_dir": evidence.output_dir, "failures": evidence.failures, "checks": evidence.checks}))
    return 1 if evidence.failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
