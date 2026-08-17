import argparse
from pathlib import Path

import cv2
import fitz
import numpy as np
from PIL import Image, ImageDraw


def _write_qr(path: Path, value: str) -> None:
    encoder = cv2.QRCodeEncoder_create()
    qr = encoder.encode(value).astype("uint8")
    qr = cv2.copyMakeBorder(qr, 20, 20, 20, 20, cv2.BORDER_CONSTANT, value=255)
    qr = cv2.resize(qr, None, fx=8, fy=8, interpolation=cv2.INTER_NEAREST)
    cv2.imwrite(str(path), qr)


def _write_barcode(path: Path, value12: str) -> None:
    l_codes = {
        "0": "0001101",
        "1": "0011001",
        "2": "0010011",
        "3": "0111101",
        "4": "0100011",
        "5": "0110001",
        "6": "0101111",
        "7": "0111011",
        "8": "0110111",
        "9": "0001011",
    }
    g_codes = {
        "0": "0100111",
        "1": "0110011",
        "2": "0011011",
        "3": "0100001",
        "4": "0011101",
        "5": "0111001",
        "6": "0000101",
        "7": "0010001",
        "8": "0001001",
        "9": "0010111",
    }
    r_codes = {
        "0": "1110010",
        "1": "1100110",
        "2": "1101100",
        "3": "1000010",
        "4": "1011100",
        "5": "1001110",
        "6": "1010000",
        "7": "1000100",
        "8": "1001000",
        "9": "1110100",
    }
    parity_patterns = {
        "0": "LLLLLL",
        "1": "LLGLGG",
        "2": "LLGGLG",
        "3": "LLGGGL",
        "4": "LGLLGG",
        "5": "LGGLLG",
        "6": "LGGGLL",
        "7": "LGLGLG",
        "8": "LGLGGL",
        "9": "LGGLGL",
    }

    def checksum12(raw: str) -> str:
        total = 0
        for index, ch in enumerate(raw, start=1):
            digit = int(ch)
            total += digit if index % 2 == 1 else 3 * digit
        return str((10 - (total % 10)) % 10)

    full = value12 + checksum12(value12)
    bits = "101"
    for digit, parity in zip(full[1:7], parity_patterns[full[0]], strict=True):
        bits += l_codes[digit] if parity == "L" else g_codes[digit]
    bits += "01010"
    for digit in full[7:]:
        bits += r_codes[digit]
    bits += "101"

    module = 4
    quiet_zone = 12
    width = (len(bits) + quiet_zone * 2) * module
    height = 160
    image = Image.new("L", (width, height), color=255)
    draw = ImageDraw.Draw(image)
    cursor = quiet_zone * module
    for bit in bits:
        if bit == "1":
            draw.rectangle([cursor, 0, cursor + module - 1, height - 1], fill=0)
        cursor += module
    image.save(path)


def _write_chart(path: Path) -> None:
    image = Image.new("RGB", (256, 160), color=(248, 250, 252))
    draw = ImageDraw.Draw(image)
    draw.rectangle([24, 20, 232, 132], outline=(148, 163, 184), width=2)
    draw.line([40, 120, 40, 36, 216, 36], fill=(148, 163, 184), width=2)
    bars = [(68, 92), (104, 72), (140, 54), (176, 46)]
    for x, top in bars:
        draw.rectangle([x, top, x + 18, 120], fill=(59, 130, 246))
    image.save(path)


def _apply_scan_noise(image: Image.Image, *, seed: int, jpeg_quality: int = 42) -> Image.Image:
    arr = np.array(image.convert("RGB"))
    blurred = cv2.GaussianBlur(arr, (5, 5), 0)
    rng = np.random.default_rng(seed)
    noise = rng.normal(loc=0.0, scale=11.0, size=blurred.shape)
    noisy = np.clip(blurred.astype("float32") + noise, 0, 255).astype("uint8")
    encoded_ok, encoded = cv2.imencode(
        ".jpg",
        cv2.cvtColor(noisy, cv2.COLOR_RGB2BGR),
        [int(cv2.IMWRITE_JPEG_QUALITY), int(jpeg_quality)],
    )
    if encoded_ok:
        decoded = cv2.imdecode(encoded, cv2.IMREAD_COLOR)
        if decoded is not None:
            noisy = cv2.cvtColor(decoded, cv2.COLOR_BGR2RGB)
    return Image.fromarray(noisy)


def _write_line_chart(path: Path) -> None:
    image = Image.new("RGB", (320, 220), color=(248, 250, 252))
    draw = ImageDraw.Draw(image)
    frame = (28, 20, 296, 194)
    draw.rectangle(frame, outline=(203, 213, 225), width=2)
    draw.text((40, 28), "Revenue trend", fill=(51, 65, 85))
    draw.line((52, 48, 52, 172), fill=(100, 116, 139), width=3)
    draw.line((52, 172, 272, 172), fill=(100, 116, 139), width=3)
    points = [(68, 152), (106, 138), (146, 144), (188, 110), (228, 94), (264, 74)]
    draw.line(points, fill=(37, 99, 235), width=5)
    for x, y in points:
        draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=(37, 99, 235))
    draw.text((48, 182), "Q1", fill=(71, 85, 105))
    draw.text((108, 182), "Q2", fill=(71, 85, 105))
    draw.text((168, 182), "Q3", fill=(71, 85, 105))
    draw.text((228, 182), "Q4", fill=(71, 85, 105))
    _apply_scan_noise(image, seed=7).save(path)


def _write_diagram(path: Path) -> None:
    image = Image.new("RGB", (256, 160), color=(250, 250, 249))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle([28, 24, 108, 60], radius=8, outline=(71, 85, 105), width=2)
    draw.rounded_rectangle([148, 24, 228, 60], radius=8, outline=(71, 85, 105), width=2)
    draw.rounded_rectangle([88, 96, 168, 132], radius=8, outline=(71, 85, 105), width=2)
    draw.line([108, 42, 148, 42], fill=(71, 85, 105), width=2)
    draw.line([128, 60, 128, 96], fill=(71, 85, 105), width=2)
    draw.line([188, 60, 128, 96], fill=(71, 85, 105), width=2)
    draw.line([68, 60, 128, 96], fill=(71, 85, 105), width=2)
    _apply_scan_noise(image, seed=11, jpeg_quality=46).save(path)


def _draw_text_block(draw: ImageDraw.ImageDraw, *, x: int, y: int, lines: list[str], fill: tuple[int, int, int]) -> int:
    cursor = y
    for line in lines:
        draw.text((x, cursor), line, fill=fill)
        cursor += 22
    return cursor


def _write_table_page(
    path: Path,
    *,
    header: list[str],
    rows: list[list[str]],
    title: str = "",
    merged_header: str = "",
    borderless: bool = False,
    leading_paragraph: list[str] | None = None,
) -> None:
    image = Image.new("RGB", (900, 1200), color=(250, 250, 249))
    draw = ImageDraw.Draw(image)
    ink = (51, 65, 85)
    accent = (100, 116, 139)
    left = 72
    top = 72
    width = 756
    row_height = 58
    col_width = width // max(1, len(header))
    cursor_y = top

    if title:
        draw.text((left, cursor_y), title, fill=ink)
        cursor_y += 36
    if leading_paragraph:
        cursor_y = _draw_text_block(draw, x=left, y=cursor_y, lines=list(leading_paragraph), fill=ink) + 20

    table_top = cursor_y
    if merged_header:
        draw.rectangle([left, cursor_y, left + width, cursor_y + row_height], outline=accent, width=2)
        draw.text((left + 16, cursor_y + 16), merged_header, fill=ink)
        cursor_y += row_height

    if borderless:
        draw.text((left, cursor_y), "   ".join(header), fill=ink)
        cursor_y += row_height
        draw.line([left, cursor_y - 10, left + width, cursor_y - 10], fill=accent, width=2)
        for row in rows:
            draw.text((left, cursor_y), "   ".join(row), fill=ink)
            cursor_y += row_height
    else:
        total_rows = 1 + len(rows)
        draw.rectangle([left, cursor_y, left + width, cursor_y + row_height * total_rows], outline=accent, width=2)
        for index in range(1, len(header)):
            x = left + index * col_width
            draw.line([x, cursor_y, x, cursor_y + row_height * total_rows], fill=accent, width=2)
        for row_index in range(1, total_rows):
            y = cursor_y + row_index * row_height
            draw.line([left, y, left + width, y], fill=accent, width=2)
        for col_index, cell in enumerate(header):
            draw.text((left + col_index * col_width + 12, cursor_y + 16), cell, fill=ink)
        for row_index, row in enumerate(rows, start=1):
            for col_index, cell in enumerate(row):
                draw.text((left + col_index * col_width + 12, cursor_y + row_index * row_height + 16), cell, fill=ink)
        cursor_y += row_height * total_rows

    draw.text((left, cursor_y + 28), f"Rows: {len(rows)}", fill=accent)
    if table_top > top:
        draw.line([left, table_top - 18, left + width, table_top - 18], fill=(226, 232, 240), width=1)
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def _write_pdf_from_image(image_path: Path, pdf_path: Path) -> None:
    with Image.open(image_path) as image:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        image.convert("RGB").save(pdf_path, "PDF")


def _write_pdf_from_images(image_paths: list[Path], pdf_path: Path) -> None:
    if not image_paths:
        return
    loaded: list[Image.Image] = []
    try:
        for path in image_paths:
            loaded.append(Image.open(path).convert("RGB"))
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        head, tail = loaded[0], loaded[1:]
        head.save(pdf_path, "PDF", save_all=True, append_images=tail)
    finally:
        for image in loaded:
            image.close()


def _write_text_pdf(page_lines: list[list[str]], pdf_path: Path) -> None:
    doc = fitz.open()
    try:
        for lines in page_lines:
            page = doc.new_page(width=612, height=792)
            page.insert_textbox(
                fitz.Rect(48, 48, 564, 744),
                "\n".join(lines),
                fontname="courier",
                fontsize=12,
                lineheight=1.4,
            )
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(pdf_path)
    finally:
        doc.close()


def _write_two_column_pdf(pdf_path: Path) -> None:
    doc = fitz.open()
    try:
        page = doc.new_page(width=612, height=792)
        page.insert_textbox(
            fitz.Rect(48, 60, 276, 744),
            "\n\n".join(
                [
                    "North region revenue increased steadily.",
                    "Operating margin held above target.",
                    "Backlog remained within forecast.",
                ]
            ),
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        page.insert_textbox(
            fitz.Rect(336, 60, 564, 744),
            "\n\n".join(
                [
                    "East region revenue accelerated in Q3.",
                    "Customer churn declined year over year.",
                    "Logistics cost stayed flat.",
                ]
            ),
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(pdf_path)
    finally:
        doc.close()


def _write_header_footer_noise_pdf(pdf_path: Path) -> None:
    doc = fitz.open()
    try:
        page = doc.new_page(width=612, height=792)
        page.insert_textbox(
            fitz.Rect(48, 24, 564, 50),
            "Quarterly Operations Report | Internal Use Only | Finance and Operations",
            fontname="courier",
            fontsize=10,
            lineheight=1.2,
        )
        page.insert_textbox(
            fitz.Rect(48, 80, 276, 744),
            "\n\n".join(
                [
                    "North region revenue increased steadily.",
                    "Operating margin held above target.",
                ]
            ),
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        page.insert_textbox(
            fitz.Rect(336, 80, 564, 744),
            "\n\n".join(
                [
                    "East region revenue accelerated in Q3.",
                    "Customer churn declined year over year.",
                ]
            ),
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        page.insert_textbox(
            fitz.Rect(120, 756, 492, 784),
            "Page 1 of 1 | Prepared for the quarterly business review",
            fontname="courier",
            fontsize=10,
            lineheight=1.2,
        )
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(pdf_path)
    finally:
        doc.close()


def _write_mixed_layout_pdf(pdf_path: Path) -> None:
    doc = fitz.open()
    try:
        page = doc.new_page(width=612, height=792)
        page.insert_textbox(
            fitz.Rect(48, 60, 276, 220),
            "North region revenue increased steadily.\n\nOperating margin held above target.",
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        page.insert_textbox(
            fitz.Rect(336, 60, 564, 220),
            "East region revenue accelerated in Q3.\n\nCustomer churn declined year over year.",
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        page.insert_textbox(
            fitz.Rect(48, 320, 564, 420),
            (
                "Diagram summary remains aligned with the layout flow and should appear "
                "after both text columns.\n\n"
                "The final synthesis block should remain below the two-column content "
                "instead of being pulled into the left column."
            ),
            fontname="courier",
            fontsize=12,
            lineheight=1.4,
        )
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(pdf_path)
    finally:
        doc.close()


def _resolve_multilingual_font() -> Path | None:
    for candidate in (
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/arphic/uming.ttc",
        "/usr/share/fonts/truetype/arphic/ukai.ttc",
    ):
        path = Path(candidate)
        if path.exists():
            return path
    return None


def _resolve_handwriting_font() -> Path | None:
    for candidate in (
        "/usr/share/fonts/truetype/arphic/ukai.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ):
        path = Path(candidate)
        if path.exists():
            return path
    return None


def _write_handwriting_note(path: Path) -> None:
    _write_handwriting_note_variant(path, text="Approved72", variant="scan")


def _write_handwriting_note_variant(path: Path, *, text: str, variant: str) -> None:
    from PIL import ImageFont

    image = Image.new("L", (460, 190), color=242)
    draw = ImageDraw.Draw(image)
    font_path = _resolve_handwriting_font()
    font = ImageFont.load_default()
    if font_path is not None:
        try:
            font = ImageFont.truetype(str(font_path), 34)
        except Exception:
            font = ImageFont.load_default()
    draw.text((30, 70), text, fill=18, font=font)
    if variant == "scan":
        image = image.rotate(-3, expand=False, fillcolor=242)
        noisy = _apply_scan_noise(image.convert("RGB"), seed=19, jpeg_quality=48).convert("L")
    elif variant == "combo":
        arr = np.array(image)
        h, w = arr.shape
        src = np.float32([[18, 18], [w - 18, 8], [24, h - 22], [w - 18, h - 12]])
        dst = np.float32([[0, 0], [w, 18], [8, h], [w, h - 14]])
        mat = cv2.getPerspectiveTransform(src, dst)
        warped = cv2.warpPerspective(arr, mat, (w, h), borderValue=242)
        img2 = Image.fromarray(warped).rotate(-5, expand=False, fillcolor=242)
        noisy = _apply_scan_noise(img2.convert("RGB"), seed=23, jpeg_quality=42).convert("L")
    else:
        noisy = image
    path.parent.mkdir(parents=True, exist_ok=True)
    noisy.save(path)


def _write_mixed_scan_memo(path: Path) -> None:
    from PIL import ImageFont

    image = Image.new("L", (920, 1180), color=241)
    draw = ImageDraw.Draw(image)
    font_path = _resolve_handwriting_font()
    title_font = ImageFont.load_default()
    body_font = ImageFont.load_default()
    if font_path is not None:
        try:
            title_font = ImageFont.truetype(str(font_path), 34)
            body_font = ImageFont.truetype(str(font_path), 28)
        except Exception:
            title_font = ImageFont.load_default()
            body_font = ImageFont.load_default()

    draw.text((86, 82), "Mixed scan memo", fill=28, font=title_font)
    left_lines = [
        "North cluster backlog stayed within forecast.",
        "West cluster returns dropped to 3 percent.",
    ]
    right_lines = [
        "APAC pilot approval moved to wave 2.",
        "Customer retention stayed at 94 percent.",
    ]
    footer_lines = [
        "Mixed scan synthesis should appear after both columns.",
        "Escalate the Jakarta handoff on Tuesday.",
    ]

    left_y = 182
    for line in left_lines:
        draw.text((88, left_y), line, fill=24, font=body_font)
        left_y += 58

    right_y = 182
    for line in right_lines:
        draw.text((494, right_y), line, fill=24, font=body_font)
        right_y += 58

    footer_y = 498
    for line in footer_lines:
        draw.text((116, footer_y), line, fill=30, font=body_font)
        footer_y += 60

    draw.line((78, 148, 848, 148), fill=168, width=2)
    draw.line((454, 162, 454, 336), fill=196, width=1)
    draw.line((96, 454, 838, 454), fill=182, width=1)

    arr = np.array(image)
    blurred = cv2.GaussianBlur(arr, (5, 5), 0)
    noise = np.random.default_rng(29).normal(loc=0.0, scale=10.5, size=blurred.shape)
    noisy = np.clip(blurred.astype("float32") + noise, 0, 255).astype("uint8")
    src = np.float32(
        [[16, 20], [arr.shape[1] - 18, 8], [28, arr.shape[0] - 18], [arr.shape[1] - 24, arr.shape[0] - 12]]
    )
    dst = np.float32([[0, 0], [arr.shape[1], 14], [8, arr.shape[0]], [arr.shape[1], arr.shape[0] - 18]])
    mat = cv2.getPerspectiveTransform(src, dst)
    warped = cv2.warpPerspective(noisy, mat, (arr.shape[1], arr.shape[0]), borderValue=241)
    final = _apply_scan_noise(Image.fromarray(warped).convert("RGB"), seed=31, jpeg_quality=43).convert("L")
    path.parent.mkdir(parents=True, exist_ok=True)
    final.save(path)


def _write_word_project_brief_docx(path: Path) -> None:
    from docx import Document as DocxDocument  # type: ignore

    doc = DocxDocument()
    doc.add_heading("Word Project Brief", level=1)
    doc.add_paragraph("Owner: Lina Chen")
    doc.add_paragraph("Delivery milestone ships on Monday.")
    try:
        doc.add_paragraph("Review the onboarding packet.", style="List Bullet")
        doc.add_paragraph("Confirm the budget note.", style="List Bullet")
    except Exception:
        doc.add_paragraph("- Review the onboarding packet.")
        doc.add_paragraph("- Confirm the budget note.")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _write_watermark_heavy_pdf(pdf_path: Path) -> None:
    _write_text_pdf(
        [
            [
                "DRAFT",
                "Company Confidential",
                "仅供内部使用",
                "",
                "Watermark-Heavy Memo",
                "",
                "Owner: Mei Lin",
                "Launch rehearsal: 2026-05-03",
                "Checklist status: ready for review.",
            ]
        ],
        pdf_path,
    )


def _write_excel_budget_sheet(path: Path) -> None:
    from openpyxl import Workbook  # type: ignore

    wb = Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws.append(["Region", "Budget", "Status"])
    ws.append(["North", 120, "Locked"])
    ws.append(["APAC", 138, "Review"])
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))


def _write_watermark_overlay_scan(path: Path) -> None:
    from PIL import ImageFont

    image = Image.new("RGB", (860, 1180), color=(244, 244, 240))
    draw = ImageDraw.Draw(image)
    font_path = _resolve_handwriting_font()
    title_font = ImageFont.load_default()
    body_font = ImageFont.load_default()
    overlay_font = ImageFont.load_default()
    if font_path is not None:
        try:
            title_font = ImageFont.truetype(str(font_path), 32)
            body_font = ImageFont.truetype(str(font_path), 26)
            overlay_font = ImageFont.truetype(str(font_path), 48)
        except Exception:
            title_font = ImageFont.load_default()
            body_font = ImageFont.load_default()
            overlay_font = ImageFont.load_default()

    draw.text((96, 96), "Watermark Overlay Scan", fill=(45, 45, 45), font=title_font)
    draw.text((96, 172), "Owner: Han Xu", fill=(35, 35, 35), font=body_font)
    draw.text((96, 226), "Escalation date: 2026-06-12", fill=(35, 35, 35), font=body_font)
    draw.text((96, 280), "Follow-up route: APAC controls team.", fill=(35, 35, 35), font=body_font)
    draw.text((96, 334), "This scan includes a heavy diagonal watermark.", fill=(35, 35, 35), font=body_font)

    overlay = Image.new("RGBA", image.size, (255, 255, 255, 0))
    odraw = ImageDraw.Draw(overlay)
    odraw.text((170, 500), "CONFIDENTIAL", fill=(140, 140, 140, 88), font=overlay_font)
    overlay = overlay.rotate(-28, resample=Image.Resampling.BICUBIC, expand=False)
    composed = Image.alpha_composite(image.convert("RGBA"), overlay).convert("RGB")
    noisy = _apply_scan_noise(composed, seed=37, jpeg_quality=40)
    path.parent.mkdir(parents=True, exist_ok=True)
    noisy.save(path)


def _write_multilingual_pdf(pdf_path: Path) -> None:
    doc = fitz.open()
    try:
        page = doc.new_page(width=612, height=792)
        font_path = _resolve_multilingual_font()
        font_name = "courier"
        if font_path is not None:
            font_name = "multilingual_cjk"
            page.insert_font(fontname=font_name, fontfile=str(font_path))
        page.insert_textbox(
            fitz.Rect(48, 48, 564, 744),
            "\n".join(
                [
                    "Multilingual revenue summary.",
                    "",
                    "APAC revenue 同比增长 12%。",
                    "North America customer retention remained 94%.",
                    "EMEA pipeline status 保持 stable。",
                    "Support contact alias is bilingual-helpdesk.",
                ]
            ),
            fontname=font_name,
            fontsize=12,
            lineheight=1.5,
        )
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(pdf_path)
    finally:
        doc.close()


def generate_assets(output_root: Path) -> None:
    targets = [
        (output_root / "qr_sheet" / "input" / "qr.png", lambda path: _write_qr(path, "HELLO-QR")),
        (output_root / "qr_sheet" / "golden" / "qr.png", lambda path: _write_qr(path, "HELLO-QR")),
        (output_root / "barcode_label" / "input" / "barcode.png", lambda path: _write_barcode(path, "590123412345")),
        (output_root / "barcode_label" / "golden" / "barcode.png", lambda path: _write_barcode(path, "590123412345")),
        (output_root / "table_scan" / "input" / "chart.png", _write_chart),
        (output_root / "table_scan" / "golden" / "chart.png", _write_chart),
        (output_root / "diagram_page" / "input" / "diagram.png", _write_diagram),
        (output_root / "diagram_page" / "golden" / "diagram.png", _write_diagram),
    ]
    for path, writer in targets:
        path.parent.mkdir(parents=True, exist_ok=True)
        writer(path)


def generate_broader_assets(output_root: Path) -> None:
    chart_image = output_root / "chart_pdf" / "input" / "chart.png"
    line_chart_image = output_root / "line_chart_pdf" / "input" / "line_chart.png"
    diagram_image = output_root / "diagram_pdf" / "input" / "diagram.png"
    qr_image = output_root / "qr_image" / "input" / "sample.png"
    barcode_image = output_root / "barcode_image" / "input" / "sample.png"
    cross_page_page_1 = output_root / "cross_page_table_pdf" / "input" / "page-1.png"
    cross_page_page_2 = output_root / "cross_page_table_pdf" / "input" / "page-2.png"
    borderless_table_image = output_root / "borderless_table_scan" / "input" / "sample.png"
    merged_header_image = output_root / "merged_header_table_pdf" / "input" / "table.png"
    leading_paragraph_image = output_root / "table_with_leading_paragraph_pdf" / "input" / "page.png"
    two_column_pdf = output_root / "two_column_pdf" / "input" / "sample.pdf"
    header_footer_noise_pdf = output_root / "header_footer_noise_pdf" / "input" / "sample.pdf"
    mixed_layout_pdf = output_root / "mixed_layout_pdf" / "input" / "sample.pdf"
    multilingual_pdf = output_root / "multilingual_pdf" / "input" / "sample.pdf"
    handwriting_note_image = output_root / "handwriting_note_image" / "input" / "sample.png"
    handwriting_skewed_note_image = output_root / "handwriting_skewed_note_image" / "input" / "sample.png"
    mixed_scan_memo_image = output_root / "mixed_scan_memo_image" / "input" / "sample.png"
    word_project_brief_docx = output_root / "word_project_brief_docx" / "input" / "sample.docx"
    watermark_heavy_pdf = output_root / "watermark_heavy_pdf" / "input" / "sample.pdf"
    excel_budget_sheet_xlsx = output_root / "excel_budget_sheet_xlsx" / "input" / "sample.xlsx"
    watermark_overlay_scan_image = output_root / "watermark_overlay_scan_image" / "input" / "sample.png"

    for path, writer in (
        (chart_image, _write_chart),
        (line_chart_image, _write_line_chart),
        (diagram_image, _write_diagram),
        (qr_image, lambda target: _write_qr(target, "HELLO-QR")),
        (barcode_image, lambda target: _write_barcode(target, "590123412345")),
    ):
        path.parent.mkdir(parents=True, exist_ok=True)
        writer(path)

    _write_table_page(
        cross_page_page_1,
        title="Quarterly revenue by region",
        header=["Region", "Q1", "Q2"],
        rows=[["North", "120", "132"], ["South", "98", "110"], ["West", "115", "121"]],
    )
    _write_table_page(
        cross_page_page_2,
        header=["Region", "Q1", "Q2"],
        rows=[["East", "107", "116"], ["Central", "111", "119"], ["APAC", "126", "138"]],
    )
    _write_table_page(
        borderless_table_image,
        title="Inventory snapshot",
        header=["Item", "Qty", "Warehouse"],
        rows=[["Paper", "220", "HZ-A"], ["Pens", "540", "HZ-B"], ["Folders", "180", "HZ-A"]],
        borderless=True,
    )
    _write_table_page(
        merged_header_image,
        title="Project budget summary",
        merged_header="Budget 2026",
        header=["Team", "Approved", "Spent"],
        rows=[["Platform", "320", "188"], ["Search", "280", "154"], ["Ops", "160", "97"]],
    )
    _write_table_page(
        leading_paragraph_image,
        title="Operations report",
        leading_paragraph=[
            "The following table summarizes the latest quarterly on-time delivery metrics.",
            "All values are percentages and should be indexed with the table content.",
        ],
        header=["Quarter", "On-time", "Delayed"],
        rows=[["Q1", "96%", "4%"], ["Q2", "94%", "6%"], ["Q3", "97%", "3%"]],
    )

    _write_pdf_from_image(chart_image, output_root / "chart_pdf" / "input" / "sample.pdf")
    _write_pdf_from_image(line_chart_image, output_root / "line_chart_pdf" / "input" / "sample.pdf")
    _write_pdf_from_image(diagram_image, output_root / "diagram_pdf" / "input" / "sample.pdf")
    _write_text_pdf(
        [
            [
                "Quarterly revenue by region.",
                "",
                "| Region | Q1 | Q2 |",
                "| --- | --- | --- |",
                "| North | 120 | 132 |",
                "| South | 98 | 110 |",
                "| West | 115 | 121 |",
            ],
            [
                "| Region | Q1 | Q2 |",
                "| --- | --- | --- |",
                "| East | 107 | 116 |",
                "| Central | 111 | 119 |",
                "| APAC | 126 | 138 |",
            ],
        ],
        output_root / "cross_page_table_pdf" / "input" / "sample.pdf",
    )
    _write_text_pdf(
        [
            [
                "Project budget summary.",
                "",
                "Budget 2026",
                "",
                "| Team | Approved | Spent |",
                "| --- | --- | --- |",
                "| Platform | 320 | 188 |",
                "| Search | 280 | 154 |",
                "| Ops | 160 | 97 |",
            ]
        ],
        output_root / "merged_header_table_pdf" / "input" / "sample.pdf",
    )
    _write_text_pdf(
        [
            [
                "The following table summarizes the latest quarterly on-time delivery metrics.",
                "",
                "All values are percentages and should be indexed with the table content.",
                "",
                "| Quarter | On-time | Delayed |",
                "| --- | --- | --- |",
                "| Q1 | 96% | 4% |",
                "| Q2 | 94% | 6% |",
                "| Q3 | 97% | 3% |",
            ]
        ],
        output_root / "table_with_leading_paragraph_pdf" / "input" / "sample.pdf",
    )
    _write_two_column_pdf(two_column_pdf)
    _write_header_footer_noise_pdf(header_footer_noise_pdf)
    _write_mixed_layout_pdf(mixed_layout_pdf)
    _write_multilingual_pdf(multilingual_pdf)
    _write_handwriting_note(handwriting_note_image)
    _write_handwriting_note_variant(handwriting_skewed_note_image, text="APAC128", variant="combo")
    _write_mixed_scan_memo(mixed_scan_memo_image)
    _write_word_project_brief_docx(word_project_brief_docx)
    _write_watermark_heavy_pdf(watermark_heavy_pdf)
    _write_excel_budget_sheet(excel_budget_sheet_xlsx)
    _write_watermark_overlay_scan(watermark_overlay_scan_image)

    for src, dst in (
        (chart_image, output_root / "chart_pdf" / "golden" / "chart.png"),
        (line_chart_image, output_root / "line_chart_pdf" / "golden" / "line_chart.png"),
        (diagram_image, output_root / "diagram_pdf" / "golden" / "diagram.png"),
        (qr_image, output_root / "qr_image" / "golden" / "sample.png"),
        (barcode_image, output_root / "barcode_image" / "golden" / "sample.png"),
        (cross_page_page_1, output_root / "cross_page_table_pdf" / "golden" / "page-1.png"),
        (cross_page_page_2, output_root / "cross_page_table_pdf" / "golden" / "page-2.png"),
        (borderless_table_image, output_root / "borderless_table_scan" / "golden" / "sample.png"),
        (merged_header_image, output_root / "merged_header_table_pdf" / "golden" / "table.png"),
        (leading_paragraph_image, output_root / "table_with_leading_paragraph_pdf" / "golden" / "page.png"),
        (handwriting_note_image, output_root / "handwriting_note_image" / "golden" / "sample.png"),
        (handwriting_skewed_note_image, output_root / "handwriting_skewed_note_image" / "golden" / "sample.png"),
        (mixed_scan_memo_image, output_root / "mixed_scan_memo_image" / "golden" / "sample.png"),
        (watermark_overlay_scan_image, output_root / "watermark_overlay_scan_image" / "golden" / "sample.png"),
    ):
        dst.parent.mkdir(parents=True, exist_ok=True)
        dst.write_bytes(src.read_bytes())


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Generate deterministic image assets for parser benchmark root fixtures."
    )
    parser.add_argument(
        "--output-root",
        default="tests/fixtures/parsing_golden",
        help="Root fixture directory to populate (default: tests/fixtures/parsing_golden)",
    )
    parser.add_argument(
        "--profile",
        choices=("smoke", "broader", "all"),
        default="smoke",
        help="Asset profile to generate (default: smoke)",
    )
    args = parser.parse_args(argv)

    output_root = Path(str(args.output_root)).resolve()
    output_root.mkdir(parents=True, exist_ok=True)
    if args.profile in {"smoke", "all"}:
        generate_assets(output_root)
    if args.profile in {"broader", "all"}:
        generate_broader_assets(output_root)
    print(f"[generate-parsing-golden-assets] wrote assets under {output_root}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
