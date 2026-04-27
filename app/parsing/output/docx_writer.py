from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX


def _add_highlighted_heading(doc: DocxDocument, title: str, *, highlight_keywords: list[str] | None) -> None:
    paragraph = doc.add_heading("", level=1)
    keywords = [str(item or "").strip() for item in (highlight_keywords or []) if str(item or "").strip()]
    if not keywords:
        paragraph.add_run(title)
        return

    cursor = 0
    while cursor < len(title):
        match = None
        for keyword in keywords:
            idx = title.find(keyword, cursor)
            if idx == -1:
                continue
            if match is None or idx < match[0]:
                match = (idx, keyword)
        if match is None:
            paragraph.add_run(title[cursor:])
            break
        idx, keyword = match
        if idx > cursor:
            paragraph.add_run(title[cursor:idx])
        run = paragraph.add_run(keyword)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        cursor = idx + len(keyword)


def write_clean_docx(
    path: str | Path,
    *,
    title: str | None,
    blocks: list[dict[str, Any]],
    highlight_keywords: list[str] | None = None,
) -> Path:
    target = Path(path)
    target.parent.mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()
    title_text = str(title or "").strip()
    if title_text:
        _add_highlighted_heading(doc, title_text, highlight_keywords=highlight_keywords)

    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        doc.add_paragraph(text)

    doc.save(str(target))
    return target


def render_clean_docx_bytes(
    *,
    title: str | None,
    blocks: list[dict[str, Any]],
    highlight_keywords: list[str] | None = None,
) -> bytes:
    doc = DocxDocument()
    title_text = str(title or "").strip()
    if title_text:
        _add_highlighted_heading(doc, title_text, highlight_keywords=highlight_keywords)

    for block in blocks or []:
        if not isinstance(block, dict):
            continue
        block_type = str(block.get("type") or "").strip().lower()
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        if block_type == "heading":
            level = int(block.get("level") or 1)
            doc.add_heading(text, level=max(1, min(level, 9)))
        else:
            doc.add_paragraph(text)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
