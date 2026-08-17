#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import re
from collections import Counter
from io import BytesIO

import pandas as pd
from docx import Document

from ..src.model import rag_tokenizer


def _table_block_type(block):
    patterns = [
        ("^(20|19)\\d{2}[年/-]\\d{1,2}[月/-]\\d{1,2}日*$", "Dt"),
        (r"^(20|19)\d{2}年$", "Dt"),
        (r"^(20|19)\d{2}[年/-]\d{1,2}月*$", "Dt"),
        ("^\\d{1,2}[月/-]\\d{1,2}日*$", "Dt"),
        (r"^第*[一二三四1-4]季度$", "Dt"),
        (r"^(20|19)\d{2}年*[一二三四1-4]季度$", "Dt"),
        (r"^(20|19)\d{2}[ABCDE]$", "DT"),
        ("^[0-9.,+%/ -]+$", "Nu"),
        (r"^[0-9A-Z/\._~-]+$", "Ca"),
        (r"^[A-Z]*[a-z' -]+$", "En"),
        (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
        (r"^.{1}$", "Sg"),
    ]
    for pattern, name in patterns:
        if re.search(pattern, block):
            return name

    tokens = [token for token in rag_tokenizer.tokenize(block).split() if len(token) > 1]
    if len(tokens) > 3:
        return "Tx" if len(tokens) < 12 else "Lx"
    if len(tokens) == 1 and rag_tokenizer.tag(tokens[0]) == "nr":
        return "Nr"
    return "Ot"


def _dominant_table_type(df):
    block_types = [
        _table_block_type(str(df.iloc[row, column]))
        for row in range(1, len(df))
        for column in range(len(df.iloc[row, :]))
    ]
    return max(Counter(block_types).items(), key=lambda item: item[1])[0]


def _header_rows(df, dominant_type):
    header_rows = [0]
    if dominant_type != "Nu":
        return header_rows
    for row in range(1, len(df)):
        row_types = Counter(_table_block_type(str(df.iloc[row, column])) for column in range(len(df.iloc[row, :])))
        if max(row_types.items(), key=lambda item: item[1])[0] != dominant_type:
            header_rows.append(row)
    return header_rows


def _active_header_offsets(header_rows, row_index):
    header_offsets = [header_row - row_index for header_row in header_rows]
    header_offsets = [offset for offset in header_offsets if offset < 0]
    index = len(header_offsets) - 1
    while index > 0:
        if header_offsets[index] - header_offsets[index - 1] > 1:
            return header_offsets[index:]
        index -= 1
    return header_offsets


def _header_prefixes(df, row_index, header_offsets):
    headers = []
    for column in range(len(df.iloc[row_index, :])):
        values = []
        for offset in header_offsets:
            value = str(df.iloc[row_index + offset, column]).strip()
            if value in values:
                continue
            values.append(value)
        header = ",".join(values)
        headers.append(f"{header}: " if header else "")
    return headers


def _compose_table_row(df, row_index, header_rows):
    header_offsets = _active_header_offsets(header_rows, row_index)
    headers = _header_prefixes(df, row_index, header_offsets)
    cells = []
    for column in range(len(df.iloc[row_index, :])):
        cell = str(df.iloc[row_index, column])
        if not cell:
            continue
        cells.append(headers[column] + cell)
    return ";".join(cells)


class IntegratedPipelineDocxParser:
    def __extract_table_content(self, tb):
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):
        if len(df) < 2:
            return []
        max_type = _dominant_table_type(df)
        colnm = len(df.iloc[0, :])
        hdrows = _header_rows(df, max_type)
        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            lines.append(_compose_table_row(df, i, hdrows))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        self.doc = Document(fnm) if isinstance(fnm, str) else Document(BytesIO(fnm))
        pn = 0  # parsed page
        secs = []  # parsed contents
        for p in self.doc.paragraphs:
            if pn > to_page:
                break

            runs_within_single_paragraph = []  # save runs within the range of pages
            for run in p.runs:
                if pn > to_page:
                    break
                if from_page <= pn < to_page and p.text.strip():
                    runs_within_single_paragraph.append(run.text)  # append run.text first

                # wrap page break checker into a static method
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1

            secs.append(
                ("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, "name") else "")
            )  # then concat run.text as part of the paragraph

        tbls = [self.__extract_table_content(tb) for tb in self.doc.tables]
        return secs, tbls


if __name__ == "__main__":
    # 1. Import required modules
    import sys

    # 2. Initialize the parser
    parser = IntegratedPipelineDocxParser()
    # 3. Provide docx file path
    docx_path = "/data/Langagent/deepdoc/data/exmaple.docx"
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]

    # 4. Call the parser
    paragraphs, tables = parser(docx_path)

    # 5. Print paragraph content
    logging.basicConfig(level=logging.INFO)
    logging.info("=== Text Paragraphs ===")
    for i, (text, style) in enumerate(paragraphs):
        if text.strip():
            logging.info("[Paragraph %s - Style: %s]: %s", i + 1, style, text)

    # 6. Print table content
    logging.info("=== Table Content ===")
    for i, table_lines in enumerate(tables):
        logging.info("[Table %s]", i + 1)
        for line in table_lines:
            logging.info("%s", line)
