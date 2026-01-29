from __future__ import annotations

from pathlib import Path

import pytest


def test_docx_fallback_parser_emits_structured_markdown(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    from docx import Document as DocxDocument  # type: ignore

    from app.parsing.parsers.docx_parser import DocxParser

    docx_path = tmp_path / "demo.docx"

    doc = DocxDocument()
    doc.add_heading("My Title", level=1)
    doc.add_paragraph("Intro paragraph.")

    # Use built-in list styles when available.
    try:
        doc.add_paragraph("Bullet item", style="List Bullet")
        doc.add_paragraph("Number item", style="List Number")
    except Exception:
        # Fallback: manual bullets still matter for governance/chunking.
        doc.add_paragraph("• Bullet item")
        doc.add_paragraph("1. Number item")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"

    doc.save(str(docx_path))

    docs = DocxParser().parse(docx_path)
    assert len(docs) == 1
    text = docs[0].page_content or ""

    assert "# My Title" in text
    assert "Intro paragraph." in text
    assert "| Name | Value |" in text
    assert "| --- | --- |" in text
