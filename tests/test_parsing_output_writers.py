from __future__ import annotations

from docx import Document as DocxDocument

from app.parsing.output.docx_writer import write_clean_docx
from app.parsing.output.markdown_writer import write_clean_markdown


def test_write_clean_markdown_persists_title_paragraphs_and_images(tmp_path) -> None:  # noqa: ANN001
    out_path = tmp_path / "clean.md"

    write_clean_markdown(
        out_path,
        title="数据采集通讯配置",
        blocks=[
            {"type": "paragraph", "text": "第一段"},
            {"type": "image", "alt": "示意图", "path": "images/a.png"},
            {"type": "paragraph", "text": "第二段"},
        ],
    )

    text = out_path.read_text(encoding="utf-8")
    assert text.startswith("# 数据采集通讯配置")
    assert "第一段" in text
    assert "![示意图](images/a.png)" in text
    assert "第二段" in text


def test_write_clean_docx_persists_heading_and_paragraphs(tmp_path) -> None:  # noqa: ANN001
    out_path = tmp_path / "clean.docx"

    write_clean_docx(
        out_path,
        title="数据采集通讯配置",
        blocks=[
            {"type": "paragraph", "text": "第一段"},
            {"type": "paragraph", "text": "第二段"},
        ],
    )

    doc = DocxDocument(str(out_path))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts[0] == "数据采集通讯配置"
    assert texts[1:] == ["第一段", "第二段"]


def test_write_clean_docx_highlights_title_keywords(tmp_path) -> None:  # noqa: ANN001
    out_path = tmp_path / "clean.docx"

    write_clean_docx(
        out_path,
        title="数据采集通讯配置",
        blocks=[],
        highlight_keywords=["通讯", "配置"],
    )

    doc = DocxDocument(str(out_path))
    heading = doc.paragraphs[0]
    highlighted = [run.text for run in heading.runs if run.font.highlight_color is not None]
    assert highlighted
    assert any("通讯" in text or "配置" in text for text in highlighted)
