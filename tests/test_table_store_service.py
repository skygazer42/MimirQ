
import re
import tempfile
import uuid
from pathlib import Path

import pytest


def test_table_store_table_id_parse_and_format():  # noqa: ANN001
    from app.services.table_store import format_table_id, parse_table_id

    doc_id = uuid.uuid4()
    tid = format_table_id(document_id=doc_id, sheet_index=3)
    assert tid.startswith("doc:")
    parsed = parse_table_id(tid)
    assert parsed is not None
    assert parsed.document_id == doc_id
    assert parsed.sheet_index == 3

    assert parse_table_id("doc:not-a-uuid:sheet:0") is None
    assert parse_table_id("doc:00000000-0000-0000-0000-000000000000:sheet:-1") is None
    assert parse_table_id("random") is None


def test_table_store_quotes_sqlite_identifiers():  # noqa: ANN001
    from app.services.table_store import quote_sqlite_ident

    assert quote_sqlite_ident("sheet_0") == '"sheet_0"'
    assert quote_sqlite_ident('sheet_"; DROP TABLE documents; --') == '"sheet_""; DROP TABLE documents; --"'


def test_table_store_import_csv_and_query(monkeypatch):  # noqa: ANN001
    from app.core.config import settings
    from app.services.table_store_service import import_table_document, run_table_query

    tenant_id = uuid.uuid4()
    dataset_id = uuid.uuid4()
    document_id = uuid.uuid4()

    with tempfile.TemporaryDirectory() as td:
        monkeypatch.setattr(settings, "TABLE_STORE_DIR", str(Path(td) / "table_store"), raising=False)

        csv_path = Path(td) / "demo.csv"
        csv_path.write_text("a,b\n1,2\n3,4\n", encoding="utf-8")

        assets = import_table_document(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            document_id=document_id,
            file_path=csv_path,
            max_rows=10,
            max_cols=10,
            sample_rows=2,
        )
        assert len(assets) == 1
        assert assets[0].table_id.startswith(f"doc:{document_id}")
        assert assets[0].row_count == 2
        assert assets[0].col_count == 2

        res = run_table_query(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            table_id=assets[0].table_id,
            sql='SELECT * FROM "sheet_0"',
            max_rows=10,
            max_cols=10,
            max_bytes=100_000,
        )
        assert res["columns"] == ["a", "b"]
        assert res["rows"][0] == [1, 2]

        # Safety: reject non-SELECT
        with pytest.raises(Exception):
            run_table_query(
                tenant_id=tenant_id,
                dataset_id=dataset_id,
                table_id=assets[0].table_id,
                sql='DELETE FROM "sheet_0"',
                max_rows=10,
                max_cols=10,
                max_bytes=100_000,
            )

        # Safety: reject multi-statement
        with pytest.raises(Exception):
            run_table_query(
                tenant_id=tenant_id,
                dataset_id=dataset_id,
                table_id=assets[0].table_id,
                sql='SELECT 1; SELECT 2',
                max_rows=10,
                max_cols=10,
                max_bytes=100_000,
            )


def test_table_store_import_xlsx_and_query(monkeypatch):  # noqa: ANN001
    import pandas as pd  # type: ignore

    from app.core.config import settings
    from app.services.table_store_service import import_table_document, run_table_query

    tenant_id = uuid.uuid4()
    dataset_id = uuid.uuid4()
    document_id = uuid.uuid4()

    with tempfile.TemporaryDirectory() as td:
        monkeypatch.setattr(settings, "TABLE_STORE_DIR", str(Path(td) / "table_store"), raising=False)

        xlsx_path = Path(td) / "demo.xlsx"
        df = pd.DataFrame({"x": [10, 20], "y": ["a", "b"]})
        df.to_excel(str(xlsx_path), index=False)

        assets = import_table_document(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            document_id=document_id,
            file_path=xlsx_path,
            max_rows=10,
            max_cols=10,
            sample_rows=1,
        )
        assert len(assets) >= 1
        table_id = assets[0].table_id

        res = run_table_query(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            table_id=table_id,
            sql='SELECT * FROM "sheet_0"',
            max_rows=10,
            max_cols=10,
            max_bytes=100_000,
        )
        assert "x" in res["columns"]
        assert len(res["rows"]) >= 1


def test_table_store_import_xlsx_with_empty_sheet(monkeypatch, tmp_path: Path):  # noqa: ANN001
    import pandas as pd  # type: ignore

    from app.core.config import settings
    from app.services.table_store_service import import_table_document, run_table_query

    tenant_id = uuid.uuid4()
    dataset_id = uuid.uuid4()
    document_id = uuid.uuid4()

    monkeypatch.setattr(settings, "TABLE_STORE_DIR", str(tmp_path / "table_store"), raising=False)
    xlsx_path = tmp_path / "empty-sheet.xlsx"
    with pd.ExcelWriter(str(xlsx_path)) as writer:
        pd.DataFrame({"name": ["a"], "phone": ["123"]}).to_excel(writer, sheet_name="电话", index=False)
        pd.DataFrame().to_excel(writer, sheet_name="Sheet1", index=False)

    assets = import_table_document(
        tenant_id=tenant_id,
        dataset_id=dataset_id,
        document_id=document_id,
        file_path=xlsx_path,
        max_rows=10,
        max_cols=10,
        sample_rows=2,
    )

    assert len(assets) == 2
    assert assets[1].sheet_name == "Sheet1"
    assert assets[1].row_count == 0
    assert assets[1].col_count == 1
    assert assets[1].columns == [{"name": "__empty__", "dtype": "object"}]

    res = run_table_query(
        tenant_id=tenant_id,
        dataset_id=dataset_id,
        table_id=assets[1].table_id,
        sql='SELECT * FROM "sheet_1"',
        max_rows=10,
        max_cols=10,
        max_bytes=100_000,
    )
    assert res["columns"] == ["__empty__"]
    assert res["rows"] == []


def test_table_store_import_docx_tables(monkeypatch):  # noqa: ANN001
    import sqlite3

    pytest.importorskip("docx")
    from docx import Document as DocxDocument  # type: ignore

    from app.core.config import settings
    from app.services.table_store import table_store_path
    from app.services.table_store_service import import_docx_tables

    tenant_id = uuid.uuid4()
    dataset_id = uuid.uuid4()
    document_id = uuid.uuid4()

    with tempfile.TemporaryDirectory() as td:
        monkeypatch.setattr(settings, "TABLE_STORE_DIR", str(Path(td) / "table_store"), raising=False)

        docx_path = Path(td) / "demo.docx"
        doc = DocxDocument()
        doc.add_paragraph("Table 1: Sales Data")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "1"
        doc.save(str(docx_path))

        assets = import_docx_tables(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            document_id=document_id,
            file_path=docx_path,
            max_rows=10,
            max_cols=10,
            sample_rows=1,
        )
        assert len(assets) == 1
        assert assets[0].sheet_name == "Table 1: Sales Data"
        assert assets[0].row_count == 1
        assert assets[0].col_count == 2

        db_path = table_store_path(tenant_id=tenant_id, dataset_id=dataset_id, document_id=document_id)
        assert db_path.exists()

        conn = sqlite3.connect(str(db_path))
        try:
            rows = conn.execute('SELECT * FROM "sheet_0"').fetchall()
        finally:
            conn.close()

        assert rows == [("A", "1")]


def test_table_store_import_db_row_snapshots(monkeypatch):  # noqa: ANN001
    from app.core.config import settings
    from app.services.table_store_service import import_db_row_snapshots, run_table_query

    tenant_id = uuid.uuid4()
    dataset_id = uuid.uuid4()
    document_id = uuid.uuid4()

    with tempfile.TemporaryDirectory() as td:
        monkeypatch.setattr(settings, "TABLE_STORE_DIR", str(Path(td) / "table_store"), raising=False)

        assets = import_db_row_snapshots(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            document_id=document_id,
            snapshots=[
                {
                    "sheet_name": "demo.users",
                    "source_table": "demo.users",
                    "source_sync_token": "tok-users-v1",
                    "columns": ["id", "name"],
                    "rows": [
                        {"id": 1, "name": "alice", "__row_pk_hash": "pkhash-1"},
                        {"id": 2, "name": "bob", "__row_pk_hash": "pkhash-2"},
                    ],
                }
            ],
            max_tables=10,
            max_rows_per_table=10,
            max_cols=10,
            sample_rows=2,
        )

        assert len(assets) == 1
        assert assets[0].sheet_name == "demo.users"
        col_names = [str(c.get("name") or "") for c in assets[0].columns]
        assert "id" in col_names
        assert "name" in col_names
        assert "__row_pk_hash" in col_names

        got = run_table_query(
            tenant_id=tenant_id,
            dataset_id=dataset_id,
            table_id=assets[0].table_id,
            sql='SELECT "id","name","__row_pk_hash" FROM "sheet_0" ORDER BY "id" ASC',
            max_rows=10,
            max_cols=10,
            max_bytes=100_000,
        )
        assert got["columns"] == ["id", "name", "__row_pk_hash"]
        assert got["rows"] == [[1, "alice", "pkhash-1"], [2, "bob", "pkhash-2"]]


def test_table_store_service_no_longer_uses_silent_pass_fallbacks() -> None:
    text = Path("app/services/table_store_service.py").read_text(encoding="utf-8")
    assert re.search(r"except[^\n]*:\n[ \t]*pass\b", text) is None
