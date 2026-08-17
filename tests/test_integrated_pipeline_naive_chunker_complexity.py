from io import BytesIO

import pytest
from docx import Document

from app.third_party.integrated_pipeline.chunkers import naive as naive_chunker


def _callback_recorder():
    calls = []

    def _callback(*args, **kwargs):
        calls.append((args, kwargs))

    return calls, _callback


def _build_docx_with_headings_and_table() -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_heading("Parent", level=1)
    document.add_heading("Child", level=2)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "A"
    document.save(buffer)
    return buffer.getvalue()


def test_by_plaintext_falls_back_to_plain_parser_when_vision_bundle_unavailable(monkeypatch):
    calls, callback = _callback_recorder()
    plain_calls = []

    class _PlainParserStub:
        def __call__(self, source, *, from_page, to_page, callback):
            plain_calls.append((source, from_page, to_page, callback))
            return [("plain section", "")], [((None, "<table></table>"), "")]

    plain_parser = _PlainParserStub()
    monkeypatch.setattr(naive_chunker, "PlainParser", lambda: plain_parser)

    def _raise_bundle(*args, **kwargs):
        raise RuntimeError("vision unavailable")

    monkeypatch.setattr(naive_chunker, "LLMBundle", _raise_bundle)

    sections, tables, parser = naive_chunker.by_plaintext(
        "sample.pdf",
        binary=b"pdf-bytes",
        callback=callback,
        layout_recognizer="Fancy Layout",
        tenant_id="tenant-1",
        lang="English",
    )

    assert sections == [("plain section", "")]
    assert tables == [((None, "<table></table>"), "")]
    assert parser is plain_parser
    assert plain_calls == [(b"pdf-bytes", 0, 100000, callback)]
    assert calls == [
        (
            (
                -1,
                "Vision layout_recognizer 'Fancy Layout' unavailable; falling back to Plain Text. (vision unavailable)",
            ),
            {},
        )
    ]


def test_markdown_extract_image_urls_with_lines_preserves_current_line_mapping():
    markdown_parser = naive_chunker.Markdown(128)
    text = '![alt](one.png)\n<p>\n  <img\n    src="two.png"\n  />\n</p>\n<img src="three.png">\n<img src="three.png">'

    assert markdown_parser.extract_image_urls_with_lines(text) == [
        {"url": "one.png", "line": 0},
        {"url": "two.png", "line": 3},
        {"url": "three.png", "line": 6},
        {"url": "three.png", "line": 7},
    ]


def test_docx_call_adds_hierarchical_table_caption_for_nearest_headings():
    lines, tables = naive_chunker.Docx()(
        filename="sample.docx",
        binary=_build_docx_with_headings_and_table(),
    )

    assert lines == [("Parent", None), ("Child", None)]
    assert tables == [
        (
            (
                None,
                "<table><caption>Table Location: sample > Parent > Child</caption>"
                "<tr><td colspan='2'>A</td></tr></table>",
            ),
            "",
        )
    ]


def test_chunk_requires_binary_for_root_embedding_extraction():
    with pytest.raises(RuntimeError, match="Embedding extraction from file path is not supported."):
        naive_chunker.chunk("sample.txt", callback=lambda *args, **kwargs: None)
